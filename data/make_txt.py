

import docx

from docx import Document
from docx.shared import Inches


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    #if not underline:
    #    u = docx.oxml.shared.OxmlElement('w:u')
    #    u.set(docx.oxml.shared.qn('w:val'), 'none')
    #    rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    #paragraph.style.font.underline = True


    return hyperlink

from docx.shared import Pt



document = Document()

# style = document.styles['Normal']
# style.font.underline=True


document.add_heading('Instructions', 0)

instructions = '''
In this task, you will read reddit posts about ebikes and extract information about the 
specific bikes and bike manufacturers on the thread. Specifically you should do the following: '''

p = document.add_paragraph(instructions)

document.add_paragraph(
    'Open the url', style='List Bullet'
)
document.add_paragraph(
    'Click to reveal all of the comments on the post (see example below)', style='List Bullet'
)

document.add_picture('comments.png', height=Inches(1))

document.add_paragraph(
    'Read all of the comments and look for any time a specific ebike manufacturer or ebike make/model is mentioned', style='List Bullet'
)

document.add_paragraph(
    'Copy and paste the exact text that references the ebike make/model or manufacturer', style='List Bullet'
)

p = document.add_paragraph(
    'If you are looking for a specific make/model, search for the make/model on the web and include a link to the specific make/model page from the manufacturers website.', style="List Bullet"
)
p = document.add_paragraph(style="List Bullet")
p.add_run(' Be sure to find the page from the manufacturer. There are lots of pages listing ebikes for sale online. The goal is the find the page from the manufacturer.').bold = True

p = document.add_paragraph(
    'In many cases, there will be many mentions of different ebikes on the same reddit thread. You should identify every single mention of an ebike make/model or manufacturer on the thread. ', style="List Bullet"
)

p = document.add_paragraph(style="List Bullet")
p.add_run(' If a make/model or manufacturer is listed twice, please include every single reference of the make/model or manufacturer. ').bold = True


document.add_paragraph(
    'The example below should help clarify this task. But please ask questions if you more confused.', style="List Bullet"
)

document.add_heading(f'Example', level=1)

example = "https://www.reddit.com/r/ebikes/comments/yaz4wv/need_an_e_bike_that_i_can_throw_2_kids_on_the/"

p = document.add_paragraph()

hyperlink = add_hyperlink(p, example, example, 'blue', False)

document.add_paragraph(
    'Start by clicking the URL. The first comment references the Riese & Muller Multicharger.', style="List Bullet"
)

document.add_picture('example1.png', height=Inches(1))

document.add_paragraph(
    'If you search for this online, you can find the manufacturer website for this make/model.', style="List Bullet"
)


hyperlink = add_hyperlink(p,
                          "https://www.r-m.de/en-us/bikes/multicharger/", 
                          "https://www.r-m.de/en-us/bikes/multicharger/",
                          'blue',
                          False)

p = document.add_paragraph(
    'Add the following link by copying and pasting the', style="List Bullet"
)
p.add_run(' exact ').bold = True

p.add_run(' text from the reddit thread and linking to the manufacturer.').bold = False

p = document.add_paragraph(style="List Bullet")

hyperlink = add_hyperlink(p, "https://www.r-m.de/en-us/bikes/multicharger/", "Riese & Muller Multicharger", 'blue', True)


p = document.add_paragraph("Now proceed to the next comment", style="List Bullet")

document.add_picture('example2.png', height=Inches(1))

p = document.add_paragraph(
    '''There is one mention of a manufacturer ('Yuba'), and one mention of a make/model ('Yuba mundo bike'). You should add them both by copying and pasting the ''', 
        style="List Bullet")

p.add_run('exact').bold = True
p.add_run(" text from the thread, like this: ")

p = document.add_paragraph(style="List Bullet 2")
p.paragraph_format.left_indent = Inches(0.25)

hyperlink = add_hyperlink(p, "https://yubabikes.com/", "Yuba", 'blue', True)

p = document.add_paragraph(style="List Bullet 2")
p.paragraph_format.left_indent = Inches(0.25)
hyperlink = add_hyperlink(p, "https://yubabikes.com/cargobikestore/mundo-electric/",
                             "Yuba mundo bike",
                             'blue', True)

p = document.add_paragraph("The first link should go to the manufacturer, because the text ('Yuba') just references a manufacturer.", style="List Bullet")

p = document.add_paragraph("The second link should go to the specific make/model 'Yuba mondo bike'. Notice that the reddit commenter does not describe the bike using the exact same string as the manufacturer, who calls the bike a 'Mundo EP8'. This is very common. You may have to use your judgment and do a little searching around online to figure out which bike you think the commenter is talking about.", style="List Bullet")

p = document.add_paragraph("The actual page you submit should look like this: ", style="List Bullet")

p = document.add_heading(f'Example', level=1)

p = document.add_paragraph(style="List Bullet")

hyperlink = add_hyperlink(p, "https://www.r-m.de/en-us/bikes/multicharger/",
                              "Riese & Muller Multicharger", 'blue', True)

p = document.add_paragraph(style="List Bullet")

hyperlink = add_hyperlink(p, "https://yubabikes.com/cargobikestore/mundo-electric/",
                             "Yuba mundo bike",
                             'blue', True)

p = document.add_paragraph(style="List Bullet")

hyperlink = add_hyperlink(p, "https://yubabikes.com/cargobikestore/mundo-electric/",
                             "Yuba mundo bike",
                             'blue', True)

p = document.add_paragraph("... plus many more examples from the remaining comments ...",
                            style="List Bullet")


with open("sample.txt", 'r') as inf:
    for ino, i in enumerate(inf):
        i = i.replace("\n", "")

        print(i)

        document.add_heading(f'{ino}', level=1)

        p = document.add_paragraph()

        hyperlink = add_hyperlink(p, i, i, 'blue', True)


        document.save('demo.docx')